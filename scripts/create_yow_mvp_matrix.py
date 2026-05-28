from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "YOW_MVP_Acceptance_Criteria_Matrix.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def style_paragraph(paragraph, size=9, bold=False, color="000000", align=None):
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.15
    if align is not None:
        paragraph.alignment = align
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.color.rgb = RGBColor.from_string(color)


def add_h(document, text, level=1):
    p = document.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(46, 116, 181) if level < 3 else RGBColor(31, 77, 120)
    return p


rows = [
    ("0. Stop the Chaos", "Master roadmap", "One canonical roadmap exists with Active, MVP Required, Post-Launch, Icebox, and Bugs sections; all scattered notes are moved or linked.", "Yes", "Missing", "Create the roadmap and make it the only planning source.", "Product"),
    ("0. Stop the Chaos", "Feature intake freeze", "No new launch features enter MVP unless they replace an existing item or fix a launch blocker.", "Yes", "Not enforced", "Add a simple decision rule and parking-lot process.", "Product"),
    ("1. Scope Freeze", "Novel-only launch", "Launch path supports novel projects first; other project types are hidden, disabled, or clearly labeled Coming Soon.", "Yes", "Partial", "Audit project type entry points and remove non-novel distractions from the launch path.", "Product/Engineering"),
    ("1. Scope Freeze", "Exact MVP list", "The MVP feature list is locked, versioned, and mapped to acceptance criteria in this matrix.", "Yes", "In progress", "Use this document as the baseline; update only through explicit decisions.", "Product"),
    ("2. Define Good Enough", "Launch blocker policy", "A blocker is limited to data loss, broken auth, broken save, broken export, unusable editor, unusable mobile layout, or missing legal/payment essentials.", "Yes", "Missing", "Adopt the blocker rule before triage begins.", "Product/QA"),
    ("3. Core Polish", "Modals and overlays", "All critical modals open, close, trap focus reasonably, preserve unsaved work, and use consistent spacing and button order.", "Conditional", "Partial", "Audit shared modal plus high-use feature modals.", "Engineering/Design"),
    ("3. Core Polish", "Save and close behavior", "Closing editors, panels, and modals never drops unsaved manuscript or worldbuilding changes; prompts appear only when useful.", "Yes", "Needs QA", "Run close/refresh/logout tests across manuscript, lore, characters, locations, and timeline.", "Engineering/QA"),
    ("3. Core Polish", "Dashboard consistency", "Dashboard lets users create, open, rename, and delete projects without layout breaks or ambiguous states.", "Yes", "Implemented, needs QA", "Test empty account, one project, many projects, and deleted-project recovery expectations.", "Engineering/QA"),
    ("3. Core Polish", "Responsive scaling", "The core writing and dashboard flows are usable on mobile and tablet widths; no key controls overlap or disappear.", "Yes", "Needs QA", "Run viewport pass for dashboard, manuscript, character, lore, timeline, and account screens.", "Design/QA"),
    ("4. Manuscript", "Stable manuscript editor", "Users can write, edit, split, and return to scenes without content loss, cursor-breaking errors, or major lag.", "Yes", "Implemented, needs QA", "Perform long-scene, rapid typing, refresh, navigation, and split-scene tests.", "Engineering/QA"),
    ("4. Manuscript", "Chapter organization", "Users can add, rename, delete, reorder, and move chapters; structure persists after refresh and login.", "Yes", "Implemented, needs QA", "Acceptance-test chapter operations and deletion warnings.", "Engineering/QA"),
    ("4. Manuscript", "Scene organization", "Users can add, rename, delete, reorder, move, and select scenes; status and notes persist.", "Yes", "Implemented, needs QA", "Test drag/drop, empty chapters, split scenes, and scene notes.", "Engineering/QA"),
    ("4. Manuscript", "Autosave", "Scene content saves within a predictable delay and survives refresh, tab close, logout/login, and network recovery where supported.", "Yes", "Implemented, high-risk QA", "Prioritize autosave reliability testing before any visual polish.", "Engineering/QA"),
    ("4. Manuscript", "Full-screen writing mode", "A distraction-reduced writing view is available and exits cleanly without losing the active scene.", "No", "Unknown/needs check", "Verify current behavior; ship only if already stable.", "Product/Engineering"),
    ("4. Manuscript", "Scroll-to-centre writing mode", "Active scene can be brought into comfortable focus without breaking navigation or layout.", "No", "Implemented partially", "Keep as polish unless it is already stable.", "Product"),
    ("4. Manuscript", "AI suggestion mode", "AI assistance can suggest prose or ideas without overwriting user text unless the user explicitly applies it.", "Conditional", "Partial", "Define the exact MVP AI action: suggest, insert, rewrite, or brainstorm.", "Product/AI"),
    ("4. Manuscript", "Passage comments/notes", "Users can attach notes to a scene or passage-level context without creating data loss or editor instability.", "No", "Partial", "De-scope passage-level annotation if scene notes are enough for launch.", "Product"),
    ("4. Manuscript", "Word count tracking", "Project, chapter, and scene word counts update accurately enough for writing progress and do not slow typing.", "Yes", "Implemented, needs QA", "Verify counts on paste, delete, scene moves, and large manuscripts.", "Engineering/QA"),
    ("5. Worldbuilding", "Character system", "Users can create, edit, delete, search, and relate characters to the active novel; data persists correctly.", "Yes", "Implemented, needs QA", "Test CRUD, empty state, many-character performance, and project scoping.", "Engineering/QA"),
    ("5. Worldbuilding", "Locations", "Users can create, edit, delete, search, and link locations to the active novel/world context.", "Yes", "Implemented, needs QA", "Test CRUD, project scoping, and links from map/timeline where applicable.", "Engineering/QA"),
    ("5. Worldbuilding", "Lore encyclopedia", "Users can create, edit, delete, categorize, search, and retrieve lore entries tied to the active project.", "Yes", "Implemented, needs QA", "Test entry lifecycle, empty states, search, and export inclusion.", "Engineering/QA"),
    ("5. Worldbuilding", "Internal linking", "Characters, locations, lore, timeline, and history can reference each other where the launch workflow requires it.", "Conditional", "Partial", "Define must-have links and postpone broad wiki-style linking.", "Product/Engineering"),
    ("5. Worldbuilding", "Search and filtering", "Users can find major project records by name/category/status without scanning long lists manually.", "Conditional", "Partial", "Audit search/filter coverage for characters, locations, lore, timeline, and ideas.", "Engineering/QA"),
    ("5. Worldbuilding", "Factions and family grouping", "Grouping supports the novelist workflow without blocking core character/lore use.", "No", "Implemented/optional", "Keep visible only if stable; otherwise move to post-launch polish.", "Product"),
    ("5. Worldbuilding", "Timeline and history linking", "Timeline events and world-history entries can be created, edited, linked/unlinked, and persist by project.", "Yes", "Implemented, needs QA", "Test linked event/history lifecycle and deletion edge cases.", "Engineering/QA"),
    ("6. Exports & Safety", "Project export", "Users can export a complete project archive containing manuscript and worldbuilding data in a restorable format.", "Yes", "Implemented, needs QA", "Test export contents for small, medium, and large projects.", "Engineering/QA"),
    ("6. Exports & Safety", "Manuscript export", "Users can export manuscript content in a readable document format with chapters/scenes in the correct order.", "Yes", "Implemented, needs QA", "Test .docx export for formatting, ordering, empty scenes, and large manuscripts.", "Engineering/QA"),
    ("6. Exports & Safety", "PDF export", "A PDF path exists, or the launch copy clearly avoids promising PDF export.", "Conditional", "Unknown/needs check", "Decide whether PDF is launch-critical or post-launch.", "Product"),
    ("6. Exports & Safety", "EPUB-ready structure", "Internal manuscript structure does not prevent future EPUB export; full EPUB generation is not required for MVP.", "No", "Preparation only", "Move full EPUB export post-launch unless already complete.", "Product/Engineering"),
    ("6. Exports & Safety", "Stylised encyclopedia export", "Not required for MVP unless it is part of the paid launch promise.", "No", "Scope trap", "Move to post-launch; keep basic project export as the launch requirement.", "Product"),
    ("6. Exports & Safety", "Backups", "Users have a clear way to avoid catastrophic data loss through cloud persistence and/or explicit export backup.", "Yes", "Partial", "Define whether backup means cloud sync, manual export, scheduled backup, or all three.", "Product/Engineering"),
    ("6. Exports & Safety", "Restore flows", "A user can restore or recover exported project data, or the product does not claim restore support yet.", "Yes", "Missing/unknown", "Create a minimal restore test before launch if exports are positioned as backups.", "Engineering/QA"),
    ("6. Exports & Safety", "Storage usage tracking", "Users can see meaningful storage usage or limits before hitting failure states.", "Conditional", "Implemented, needs QA", "Test account storage card against realistic manuscript/worldbuilding sizes.", "Engineering/QA"),
    ("7. Launch", "Authentication", "Sign up, sign in, sign out, session restore, and password/error states work reliably across refreshes.", "Yes", "Implemented, needs QA", "Run auth regression pass with fresh user, returning user, invalid credentials, and expired session.", "Engineering/QA"),
    ("7. Launch", "Autosave reliability", "Autosave has explicit test coverage or manual proof for refresh, navigation, rapid typing, and multi-scene work.", "Yes", "High-risk QA", "Make this the first launch-readiness test pass.", "Engineering/QA"),
    ("7. Launch", "Large project performance", "Dashboard, manuscript, search, and export remain usable on a realistic large novel project.", "Yes", "Needs QA", "Create/load a stress project and measure obvious failures.", "Engineering/QA"),
    ("7. Launch", "Empty/loading/error states", "Core screens give clear empty, loading, and error feedback without dead ends.", "Conditional", "Partial", "Audit top-level routes and each MVP module.", "Design/Engineering"),
    ("7. Launch", "Landing page", "Landing page explains the novelist-first value proposition and routes cleanly into signup/pricing.", "Conditional", "Implemented, needs polish", "Improve only what affects conversion and expectation-setting.", "Product/Design"),
    ("7. Launch", "SEO basics", "Title, description, social preview, canonical basics, and indexability are set for public marketing pages.", "Conditional", "Unknown/needs check", "Run a lightweight metadata audit.", "Marketing/Engineering"),
    ("7. Launch", "Analytics", "Privacy-conscious analytics capture signup, activation, project creation, and export events.", "Conditional", "Unknown/needs decision", "Decide minimum useful events before adding tooling.", "Product"),
    ("Post-Launch", "Advanced map builder", "Map features do not block novelist-first launch.", "No", "Implemented/optional", "Keep if stable; do not polish at the expense of manuscript/data safety.", "Product"),
    ("Post-Launch", "D&D project types", "Non-novel project categories are clearly deferred.", "No", "Post-launch", "Label Coming Soon or hide until novel workflow is stable.", "Product"),
    ("Post-Launch", "Collaboration", "Multi-user editing and sharing are excluded from MVP.", "No", "Post-launch", "Document as future roadmap only.", "Product"),
    ("Post-Launch", "Marketplace/community", "Marketplace and public community systems are excluded from MVP.", "No", "Post-launch", "Avoid launch copy that implies community availability.", "Product"),
    ("Post-Launch", "Mobile apps", "Native mobile apps are excluded; responsive web usability is the MVP requirement.", "No", "Post-launch", "Focus mobile effort on web responsiveness.", "Product"),
    ("Post-Launch", "Advanced AI agents", "Agentic workflows are excluded; MVP AI stays bounded and user-directed.", "No", "Post-launch", "Define one or two safe AI assistance actions for launch.", "Product/AI"),
    ("Post-Launch", "Public sharing", "Public project/profile sharing is excluded from MVP unless needed for launch marketing.", "No", "Post-launch", "Keep private workspace promise clear.", "Product"),
]


doc = Document()
section = doc.sections[0]
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width = Inches(11)
section.page_height = Inches(8.5)
for side in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
    setattr(section, side, Inches(0.55))

styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"].font.size = Pt(10)
styles["Normal"].paragraph_format.space_after = Pt(6)
styles["Normal"].paragraph_format.line_spacing = 1.25

title = doc.add_paragraph()
title.paragraph_format.space_after = Pt(2)
run = title.add_run("YOW MVP Acceptance Criteria Matrix")
run.font.name = "Calibri"
run.font.size = Pt(22)
run.font.bold = True
run.font.color.rgb = RGBColor(11, 37, 69)

subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(10)
run = subtitle.add_run("Launch triage sheet for Your Own World as a novel-first writing and worldbuilding platform.")
run.font.name = "Calibri"
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(85, 85, 85)

add_h(doc, "Launch Rule", 1)
p = doc.add_paragraph()
p.add_run("Recommended blocker rule: ").bold = True
p.add_run("only data loss, broken auth, broken save, broken export, unusable editor, unusable responsive layout, or missing legal/payment essentials should block MVP launch.")
style_paragraph(p, size=10)

add_h(doc, "Immediate Focus", 1)
for item in [
    "Lock this matrix as the MVP baseline and stop adding launch features.",
    "Run autosave/data-safety tests before visual polish.",
    "Decide which conditional items are launch promises versus post-launch polish.",
    "Use the Current Status column as a triage queue, not as final proof of readiness.",
]:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(item)
    style_paragraph(p, size=10)

add_h(doc, "Acceptance Criteria Matrix", 1)
p = doc.add_paragraph()
p.add_run("Status labels are based on the checklist plus a light codebase scan. Anything marked needs QA should be verified before being treated as launch-ready.")
style_paragraph(p, size=9, color="555555")

headers = ["Phase", "Feature", "MVP Requirement / Acceptance Criteria", "Blocker?", "Current Status", "Next Action", "Owner/Notes"]
widths = [1400, 1700, 3700, 950, 1600, 2900, 1886]
table = doc.add_table(rows=1, cols=len(headers))
table.style = "Table Grid"
for idx, header in enumerate(headers):
    cell = table.rows[0].cells[idx]
    cell.text = header
    set_cell_shading(cell, "E8EEF5")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for p in cell.paragraphs:
        style_paragraph(p, size=8, bold=True, color="0B2545", align=WD_ALIGN_PARAGRAPH.CENTER)

for row in rows:
    cells = table.add_row().cells
    for idx, value in enumerate(row):
        cells[idx].text = value
        cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if row[3] == "Yes":
            color = "FCE4E4" if idx == 3 else None
        elif row[3] == "Conditional":
            color = "FFF2CC" if idx == 3 else None
        else:
            color = "E2F0D9" if idx == 3 else None
        if color:
            set_cell_shading(cells[idx], color)
        if idx == 0 and row[0] == "Post-Launch":
            set_cell_shading(cells[idx], "F2F4F7")
        for p in cells[idx].paragraphs:
            align = WD_ALIGN_PARAGRAPH.CENTER if idx in (3, 4, 6) else None
            style_paragraph(p, size=7.5, bold=(idx == 1), color="000000", align=align)

set_table_widths(table, widths)

add_h(doc, "Recommended Sequencing", 1)
seq = [
    ("1", "Data safety and autosave", "Prove content survives refresh, navigation, logout/login, export, and large-project usage."),
    ("2", "Writing loop", "Tighten dashboard, manuscript editor, chapters, scenes, word counts, and basic manuscript export."),
    ("3", "Worldbuilding loop", "Verify characters, locations, lore, timeline/history, search/filtering, and required links."),
    ("4", "Launch polish", "Fix responsiveness, empty/loading/error states, landing page clarity, SEO basics, and analytics."),
]
seq_table = doc.add_table(rows=1, cols=3)
seq_table.style = "Table Grid"
for i, h in enumerate(["Order", "Workstream", "Definition of Done"]):
    seq_table.rows[0].cells[i].text = h
    set_cell_shading(seq_table.rows[0].cells[i], "E8EEF5")
    for p in seq_table.rows[0].cells[i].paragraphs:
        style_paragraph(p, size=9, bold=True, color="0B2545", align=WD_ALIGN_PARAGRAPH.CENTER)
for s in seq:
    cs = seq_table.add_row().cells
    for i, v in enumerate(s):
        cs[i].text = v
        for p in cs[i].paragraphs:
            style_paragraph(p, size=9, bold=(i == 1), align=WD_ALIGN_PARAGRAPH.CENTER if i == 0 else None)
set_table_widths(seq_table, [900, 2600, 9860])

doc.save(OUT)
print(OUT)
